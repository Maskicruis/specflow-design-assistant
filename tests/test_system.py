import json
import io
import time
from pathlib import Path
import pytest
import pymupdf as fitz
from docx import Document
from fastapi.testclient import TestClient
from app import store,pipeline,retrieval,models
from app.main import app

@pytest.fixture(autouse=True)
def isolated_store(tmp_path,monkeypatch):
    monkeypatch.setattr(store,'DATA',tmp_path/'data')
    monkeypatch.setattr(store,'DB',tmp_path/'data'/'test.sqlite3')
    monkeypatch.setattr(store,'SOURCES',tmp_path/'规范文件')
    store.SOURCES.mkdir()
    store.DATA.mkdir()
    for sub in ('previews','assets','imports'):
        (store.DATA/sub).mkdir()
    store.init()
    retrieval.CACHE.clear()
    yield
    retrieval.CACHE.clear()

def native_pdf(path):
    pdf=fitz.open()
    page=pdf.new_page()
    page.insert_text((50,70),'3.0.1 站址选择应根据电力系统规划、负荷分布和交通运输条件，进行技术经济比较。',fontname='china-s',fontsize=11)
    page.insert_text((50,110),'3.0.2 变电站站址应避开不良地质地段，并考虑排水和环境保护条件。',fontname='china-s',fontsize=11)
    pdf.save(path)
    return path

def scan_pdf(path):
    source=fitz.open()
    p=source.new_page()
    p.insert_text((30,60),'IMAGE ONLY',fontsize=20)
    img=p.get_pixmap().tobytes('png')
    pdf=fitz.open()
    page=pdf.new_page()
    page.insert_image(page.rect,stream=img)
    pdf.save(path)
    return path

def test_native_import_retrieval_draft_and_duplicate(tmp_path):
    path=native_pdf(tmp_path/'变电站报批稿.pdf')
    r=pipeline.import_file(path)
    assert pipeline.import_file(path)['duplicate']
    assert store.stats()['documents']==1
    hits,_=retrieval.search('变电站站址选择3.0.1')
    assert hits and hits[0]['clause']=='3.0.1'
    assert hits[0]['page']==1 and hits[0]['draft']==1
    response=retrieval.answer('变电站站址选择要求',use_model=False)
    assert response['mode']=='evidence'
    assert any('报批稿' in x for x in response['cautions'])
    assert store.one('SELECT status FROM documents WHERE id=?',(r['id'],))['status']=='ready'

def test_scanned_page_is_not_silently_indexed(tmp_path):
    r=pipeline.import_file(scan_pdf(tmp_path/'scan.pdf'))
    assert store.stats()['pending_pages']==1
    assert store.stats()['chunks']==0
    assert store.one('SELECT status FROM documents WHERE id=?',(r['id'],))['status']=='awaiting_vision'
    response=retrieval.answer('火星基地反物质引擎',use_model=False)
    assert response['mode']=='insufficient' and not response['findings']

def test_docx_native_tables_merges_and_formatting(tmp_path):
    doc=Document()
    p=doc.add_paragraph()
    p.add_run('5.1.1 变电站消防设计应复核原文。测试样本，不是工程标准。').bold=True
    table=doc.add_table(rows=3,cols=2)
    table.cell(0,0).merge(table.cell(0,1)).text='跨列表头'
    table.cell(1,0).text='名称'
    table.cell(1,1).text='单位'
    table.cell(2,0).text='演示道路'
    table.cell(2,1).text='m'
    path=tmp_path/'sample.docx';doc.save(path)
    d=pipeline.import_file(path)
    assert store.stats()['tables']==1
    table_block=store.one("SELECT detail FROM blocks WHERE kind='table'")
    detail=json.loads(table_block['detail'])
    assert 'gridSpan' in detail['ooxml'] and detail['rows'][0][0]=='跨列表头'
    para=json.loads(store.one("SELECT detail FROM blocks WHERE kind='paragraph'")['detail'])
    assert para['runs'][0]['bold'] is True
    assert store.one('SELECT kind FROM documents WHERE id=?',(d['id'],))['kind']=='docx'

def test_vlm_uncertainty_excluded_and_success_resumable(tmp_path,monkeypatch):
    d=pipeline.import_file(scan_pdf(tmp_path/'scan.pdf'))
    fake={'blocks':[{'type':'paragraph','text':'3.0.1 安全出口的布置应根据使用人数核定。','clause':'3.0.1','bbox':[0,0,1,.3]},
                    {'type':'paragraph','text':'5.0.1 防火间距为[不可辨认]米。','uncertain':True,'bbox':[0,.3,1,.5]}],
          'warnings':[],'page_label':'1','usage':{}}
    monkeypatch.setattr(pipeline,'interpret_image',lambda *a,**kw:fake)
    pipeline.interpret_page(d['id'],1)
    assert store.stats()['pending_pages']==0
    assert not retrieval.search('防火间距')[0]
    assert retrieval.search('安全出口')[0]
    count=store.stats()['chunks']
    pipeline.interpret_page(d['id'],1)
    assert store.stats()['chunks']==count

def test_invalid_citation_falls_back_to_real_evidence(tmp_path,monkeypatch):
    pipeline.import_file(native_pdf(tmp_path/'变电站.pdf'))
    store.save_settings({'enabled':True})
    monkeypatch.setattr(retrieval,'call_model',lambda *a,**k:(json.dumps({'overview':'伪造','findings':[{'text':'无来源要求','citations':['S999']}],'missing':[]}),{}))
    response=retrieval.answer('变电站站址选择')
    assert response['mode']=='evidence'
    assert all('无来源要求' not in f['text'] for f in response['findings'])
    assert any('引用校验' in x for x in response['cautions'])

def test_reindex_preserves_unchanged_bookmarks(tmp_path):
    d=pipeline.import_file(native_pdf(tmp_path/'变电站.pdf'))
    hit=retrieval.search('站址选择')[0][0]
    store.execute('INSERT INTO bookmarks VALUES (?,?)',(hit['id'],store.now()))
    pipeline.reindex_document(d['id'])
    assert len(store.rows('SELECT * FROM bookmarks'))==1
    before=store.stats()['chunks']
    pipeline.reindex_document(d['id'])
    assert store.stats()['chunks']==before

def test_api_guards_file_validation_and_paths(tmp_path):
    client=TestClient(app)
    assert client.get('/api/bootstrap').status_code==200
    assert client.post('/api/import/folder',headers={'origin':'https://evil.example'}).status_code==403
    assert client.get('/api/health',headers={'host':'evil.example'}).status_code==403
    assert client.post('/api/import/upload',files={'file':('bad.pdf',b'not a PDF','application/pdf')}).status_code==400
    assert client.post('/api/import/upload',files={'file':('old.doc',b'x','application/msword')}).status_code==400
    assert client.get('/api/documents/not-a-document/original').status_code==404
    assert client.get('/api/assets/evil.svg').status_code==400
    assert 'api_key' not in client.get('/api/settings').json()

def test_file_upload_roundtrip(tmp_path):
    client=TestClient(app)
    path=native_pdf(tmp_path/'native.pdf')
    response=client.post('/api/import/upload',files={'file':('../../safe.pdf',path.read_bytes(),'application/pdf')})
    assert response.status_code==200
    job=response.json()['job_id']
    for _ in range(100):
        status=store.one('SELECT status FROM jobs WHERE id=?',(job,))['status']
        if status not in ('queued','running'):break
        time.sleep(.05)
    assert status=='completed'
    doc=client.get('/api/documents').json()[0]
    assert store.resolve_source(doc['path']).parent==store.DATA/'imports'
    assert client.get(f'/api/documents/{doc["id"]}/original').content[:5]==b'%PDF-'
    assert client.get(f'/api/documents/{doc["id"]}/pages/1/image').headers['content-type']=='image/jpeg'

def test_model_output_validation():
    with pytest.raises(models.ModelError):models.parse_json('not-json')
    with pytest.raises(models.ModelError):models.validate_endpoint('http://remote.example/v1')
    with pytest.raises(models.ModelError):models.validate_endpoint('https://user:secret@example.com/v1')
    assert models.validate_endpoint('http://127.0.0.1:1234/v1')=='127.0.0.1'

def test_docx_image_is_preserved_and_uses_vision(tmp_path,monkeypatch):
    pdf=fitz.open();page=pdf.new_page(width=120,height=80)
    page.draw_rect(fitz.Rect(10,10,80,60),color=(0,0,1))
    image=tmp_path/'figure.png';page.get_pixmap().save(image)
    doc=Document();doc.add_paragraph('3.0.1 消防设施布置测试说明，应结合图示复核。')
    doc.add_picture(str(image));path=tmp_path/'image.docx';doc.save(path)
    result=pipeline.import_file(path)
    assert store.stats()['pending_pages']==1
    assert len(list((store.DATA/'assets').glob('*.png')))==1
    fake={'blocks':[{'type':'figure','text':'图示为一个蓝色矩形轮廓，无尺寸标注。','bbox':[0,0,1,1],'uncertain':False}], 'warnings':[]}
    monkeypatch.setattr(pipeline,'interpret_image',lambda *a,**k:fake)
    pipeline.interpret_page(result['id'],2)
    assert store.stats()['pending_pages']==0
    assert store.one("SELECT COUNT(*) n FROM chunks WHERE kind='figure'")['n']==1

def test_unreviewed_table_values_withheld_from_model(tmp_path,monkeypatch):
    result=pipeline.import_file(native_pdf(tmp_path/'变电站.pdf'))
    with store.connect() as db:
        pipeline.add_chunk(db,result['id'],1,'消防车道净宽 98765 米','table',heading='消防车道表',detail={'rows':[['净宽','98765']]})
    captured={}
    def model(messages,**kwargs):
        captured['messages']=messages
        return json.dumps({'overview':'请复核原页','findings':[],'missing':['表格需复核']}),{}
    store.save_settings({'enabled':True})
    monkeypatch.setattr(retrieval,'call_model',model)
    retrieval.answer('消防车道净宽')
    assert '98765' not in str(captured['messages'])
    assert 'table_requires_review' in captured['messages'][1]['content']

def test_cross_page_continuation_and_commentary(tmp_path):
    path=native_pdf(tmp_path/'变电站.pdf');d=pipeline.import_file(path)
    doc_id=d['id']
    with store.connect() as db:
        db.execute("UPDATE pages SET text=? WHERE doc_id=? AND number=1",('3.0.1 站址应综合考虑电网规划、环境保护、排水条件，并通过技术经济比较确定。后续条件：',doc_id))
        db.execute('INSERT INTO pages(doc_id,number,method,status,text,width,height,meta) VALUES (?,?,?,?,?,?,?,?)',(doc_id,2,'native_pdf','indexed','还应考虑交通运输、施工条件以及后续扩建的可能性，相关条件均应在设计阶段复核。\n3.0.2 新的条文内容应保留。',600,800,'{}'))
        db.execute('INSERT INTO pages(doc_id,number,method,status,text,width,height,meta) VALUES (?,?,?,?,?,?,?,?)',(doc_id,3,'native_pdf','indexed','条文说明',600,800,'{}'))
        db.execute('INSERT INTO pages(doc_id,number,method,status,text,width,height,meta) VALUES (?,?,?,?,?,?,?,?)',(doc_id,4,'native_pdf','indexed','3.0.1 本条站址选择是解释性材料，不应替代正文内容。',600,800,'{}'))
    pipeline.reindex_document(doc_id)
    retrieval.CACHE.clear()
    hits,_=retrieval.search('变电站站址3.0.1')
    first=hits[0]
    assert first['page']==1 and first['page_end']==2
    assert '交通运输' in first['text']
    assert '解释性材料' not in first['text']


def test_managed_sources_and_uploads_survive_directory_move(tmp_path,monkeypatch):
    import shutil
    source=store.SOURCES/'subfolder'
    source.mkdir()
    first=pipeline.import_file(native_pdf(source/'电力规范.pdf'))
    shutil.copy2(source/'电力规范.pdf',store.DATA/'imports'/'uploaded.pdf')
    second=pipeline.import_file(store.DATA/'imports'/'uploaded.pdf')
    # The PDFs have equal bytes: same identity, and the latest verified location wins.
    assert first['id']==second['id']
    assert store.one('SELECT path FROM documents')['path'].startswith('data:imports/')
    hit=retrieval.search('站址选择')[0][0]
    store.execute('INSERT INTO bookmarks VALUES (?,?)',(hit['id'],store.now()))
    moved=tmp_path/'new computer'
    shutil.copytree(store.DATA,moved/'data')
    shutil.copytree(store.SOURCES,moved/'规范文件')
    old_data=store.DATA
    monkeypatch.setattr(store,'DATA',moved/'data')
    monkeypatch.setattr(store,'SOURCES',moved/'规范文件')
    monkeypatch.setattr(store,'DB',store.DATA/'test.sqlite3')
    # Remove the old import to ensure the source viewer is using the new location.
    (old_data/'imports'/'uploaded.pdf').unlink()
    store.init()
    client=TestClient(app)
    assert client.get(f'/api/documents/{first["id"]}/original').content.startswith(b'%PDF-')
    assert client.get(f'/api/documents/{first["id"]}/pages/1/image').status_code==200
    assert len(client.get('/api/bookmarks').json())==1


def test_legacy_windows_path_migration_verifies_checksum(tmp_path):
    path=native_pdf(store.SOURCES/'规范.pdf')
    result=pipeline.import_file(path)
    legacy='Z:\\old-computer\\规范文件\\规范.pdf'
    store.execute('UPDATE documents SET path=? WHERE id=?',(legacy,result['id']))
    good=path.read_bytes()
    path.write_bytes(b'wrong file with same name')
    store.migrate_source_paths()
    assert store.one('SELECT path FROM documents')['path']==legacy
    path.write_bytes(good)
    store.migrate_source_paths()
    assert store.one('SELECT path FROM documents')['path']=='sources:规范.pdf'


def test_source_locator_blocks_traversal_and_missing_file_returns_404(tmp_path):
    for locator in ('sources:../private.pdf','data:imports/../../private.pdf','sources:..\\private.pdf'):
        with pytest.raises(ValueError):
            store.resolve_source(locator)
    path=native_pdf(store.SOURCES/'missing.pdf')
    result=pipeline.import_file(path)
    path.unlink()
    assert TestClient(app).get(f'/api/documents/{result["id"]}/original').status_code==404


def test_graph_removed_query_and_workflow_remain():
    client=TestClient(app)
    assert client.get('/api/graph').status_code==404
    html=client.get('/').text
    assert 'data-nav="graph"' not in html
    assert 'data-nav="workflow"' in html and 'data-nav="query"' in html
    assert client.get('/api/health').json()['version']=='1.01'


def test_portable_backup_contains_database_and_sources_but_no_env(tmp_path,monkeypatch):
    import manage
    import zipfile
    path=native_pdf(store.SOURCES/'original.pdf')
    pipeline.import_file(path)
    monkeypatch.setattr(manage,'ROOT',tmp_path)
    monkeypatch.setattr(manage,'DATA',store.DATA)
    monkeypatch.setattr(manage,'SOURCES',store.SOURCES)
    # Backup expects the production filename, unlike the rest of the fixture.
    with store.connect() as src:
        import sqlite3
        from contextlib import closing
        with closing(sqlite3.connect(store.DATA/'knowledge.sqlite3')) as dst:src.backup(dst)
    monkeypatch.setattr(manage,'health',lambda:None)
    (tmp_path/'.env').write_text('PRIVATE_TEST_KEY=do-not-copy',encoding='utf-8')
    result=manage.backup()
    with zipfile.ZipFile(result) as archive:
        assert archive.testzip() is None
        assert 'data/knowledge.sqlite3' in archive.namelist()
        assert '规范文件/original.pdf' in archive.namelist()
        assert '.env' not in archive.namelist()
    monkeypatch.setattr(manage,'health',lambda:{'app':'SpecFlow'})
    with pytest.raises(RuntimeError):manage.backup()


def test_maintenance_wont_stop_other_installation(monkeypatch):
    import manage
    monkeypatch.setattr(manage,'health',lambda:{'app':'SpecFlow','instance':'different','pid':123})
    monkeypatch.setattr(manage.os,'kill',lambda *args:pytest.fail('Must not kill another installation'))
    with pytest.raises(RuntimeError):manage.stop()
